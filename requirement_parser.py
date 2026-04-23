"""Parsers that convert Word tables into structured requirement blocks."""
from __future__ import annotations

import re
from collections import OrderedDict
from pathlib import Path
from typing import Iterator, List, Optional, Sequence, Union

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

from requirement_models import RequirementBlock, normalize_block

_ModuleHeading = re.compile(r"^\d+(?:\.\d+)*\.?\s+.+")
_HEADER_TITLES = {"功能", "模块", "序号"}
_HEADER_DESCRIPTIONS = {"描述", "内容", "说明", "需求描述"}


def parse_requirement_tables(docx_path: Union[str, Path]) -> List[RequirementBlock]:
    """Parse requirement tables from ``docx_path`` into normalized blocks.

    Falls back to paragraph-based parsing when no table blocks are found.
    """

    document = Document(Path(docx_path))
    blocks: List[RequirementBlock] = []
    current_module = "Uncategorized"
    table_index = 0

    for item in _iter_block_items(document):
        if isinstance(item, Paragraph):
            module_title = _maybe_extract_module(item.text)
            if module_title:
                current_module = module_title
            continue

        table_index += 1
        table_blocks = _parse_table(item, current_module, table_index)
        if table_blocks:
            blocks.extend(table_blocks)

    normalized = [normalize_block(block) for block in blocks if block.title and block.description]
    if normalized:
        return normalized

    # No table-based blocks found — fall back to paragraph-based parsing.
    return _parse_requirement_paragraphs(document)


def _parse_requirement_paragraphs(document) -> List[RequirementBlock]:
    """Parse a paragraph-structured (non-table) requirements document.

    Dynamically detects the minimum heading depth in the document and uses it
    as the module-context threshold.  Top-level headings become module context;
    all deeper headings each become a RequirementBlock.
    """

    def _heading_depth(text: str) -> Optional[int]:
        stripped = text.lstrip("*-# ")
        if not _ModuleHeading.match(stripped):
            return None
        section_part = stripped.split()[0].rstrip(".")
        return len(section_part.split("."))

    # First pass: find minimum heading depth to set context threshold.
    min_depth: Optional[int] = None
    for para in document.paragraphs:
        d = _heading_depth(para.text.strip())
        if d is not None and (min_depth is None or d < min_depth):
            min_depth = d
    if min_depth is None:
        return []

    blocks: List[RequirementBlock] = []
    parent_module: str = "未分类"
    current_heading: Optional[str] = None
    current_paragraphs: List[str] = []
    block_index = 0

    def _flush() -> None:
        nonlocal block_index
        if current_heading and current_paragraphs:
            description = "\n".join(current_paragraphs)
            blocks.append(
                RequirementBlock(
                    module=parent_module,
                    title=current_heading,
                    description=description,
                    notes=None,
                    block_id=f"P{block_index:03d}-001",
                    table_index=0,
                    row_index=block_index,
                    segments=_segment_text(description),
                )
            )
            block_index += 1

    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        depth = _heading_depth(text)
        if depth is not None:
            _flush()
            if depth <= min_depth:
                parent_module = text
            current_heading = text
            current_paragraphs = []
        else:
            if current_heading is not None:
                current_paragraphs.append(text)

    _flush()
    return [normalize_block(b) for b in blocks if b.title and b.description]


def _segment_text(text: str) -> Optional[List[str]]:
    if not text:
        return None

    paragraphs = [line.strip() for line in text.split("\n") if line.strip()]
    return paragraphs or None


def _make_block_id(table_index: int, row_index: int) -> str:
    return f"T{table_index:03d}-R{row_index:03d}"


def _iter_block_items(document: Document) -> Iterator[Union[Paragraph, Table]]:
    """Yield both paragraphs and tables in document order."""

    parent = document.element.body
    for child in parent.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def _maybe_extract_module(text: str) -> Optional[str]:
    stripped = text.strip()
    if not stripped:
        return None

    simplified = stripped.lstrip("*-# ")
    if not simplified:
        return None

    if _ModuleHeading.match(simplified):
        return simplified

    return None


def _parse_table(table: Table, module: str, table_index: int) -> List[RequirementBlock]:
    if not table.rows:
        return []

    header_cells = _extract_row_text(table.rows[0])
    if not _is_requirement_header(header_cells):
        return []

    parsed: List[RequirementBlock] = []
    last_block: Optional[RequirementBlock] = None

    for row_index, row in enumerate(table.rows[1:], start=1):
        col_text = _extract_row_text(row)
        col_text += [""] * (3 - len(col_text))
        title, description, example = col_text[:3]

        if not any(col_text):
            continue

        if _is_requirement_header(col_text):
            # Occasionally headers repeat mid-table after page breaks.
            continue

        if not title and last_block:
            if description:
                last_block.description = _merge_text(last_block.description, description)
                last_block.segments = _segment_text(last_block.description)
            if example:
                last_block.notes = _merge_text(last_block.notes or "", example)
            continue

        block = RequirementBlock(
            module=module,
            title=title or (last_block.title if last_block else module),
            description=description,
            notes=example or None,
            block_id=_make_block_id(table_index, row_index),
            table_index=table_index,
            row_index=row_index,
            segments=_segment_text(description),
        )
        parsed.append(block)
        last_block = block

    return parsed


def _segment_text(text: str) -> Optional[List[str]]:
    if not text:
        return None

    paragraphs = [line.strip() for line in text.split("\n") if line.strip()]
    return paragraphs or None


def _make_block_id(table_index: int, row_index: int) -> str:
    return f"T{table_index:03d}-R{row_index:03d}"


def _extract_row_text(row) -> List[str]:
    values: List[str] = []


def _extract_row_text(row) -> List[str]:
    values: List[str] = []
    seen_cells: OrderedDict[int, None] = OrderedDict()
    for cell in row.cells:
        cell_id = id(cell._tc)
        if cell_id in seen_cells:
            continue
        seen_cells[cell_id] = None
        values.append(_clean_cell_text(cell))
    return values


def _clean_cell_text(cell: _Cell) -> str:
    if not cell.text:
        return ""

    lines = [paragraph.text.strip() for paragraph in cell.paragraphs]
    filtered = [line for line in lines if line]
    return "\n".join(filtered).strip()


def _is_requirement_header(values: Sequence[str]) -> bool:
    if not values:
        return False

    clean = [value.replace("：", "").strip() for value in values if value]
    if not clean:
        return False

    first = clean[0]
    second = clean[1] if len(clean) > 1 else ""
    return first in _HEADER_TITLES and second in _HEADER_DESCRIPTIONS


def _merge_text(existing: str, addition: str) -> str:
    existing = existing.strip()
    addition = addition.strip()
    if not existing:
        return addition
    if not addition:
        return existing
    return f"{existing}\n{addition}"
