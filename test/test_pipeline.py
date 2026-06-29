import asyncio
import io
import sys
from pathlib import Path
from types import SimpleNamespace

import pytest
from docx import Document

from generate_requirement_prompts import LLMOutputParser
from generate_requirement_prompts import validate_docx
from pipeline.config import BatchConfig, ExecutionFlags, LLMConfig, PathsConfig, PipelineConfig
from pipeline.llm_client import LLMClient, LLMRequest, LLMResponse
from pipeline.llm_providers.base import ProviderResult
from pipeline.llm_providers.custom_http import CustomHTTPProvider
from pipeline.llm_providers.factory import build_provider
from pipeline.llm_providers.openai_compatible import OpenAICompatibleProvider
from requirement_models import RequirementBlock
from pipeline.state import PipelineState
from pipeline.stages import PipelineStages
from pipeline.stages import run_pipeline


def test_pipeline_config_from_dict(tmp_path: Path) -> None:
    payload = {
        "paths": {
            "docx": "custom.docx",
            "artifacts_dir": "artifacts",
            "prompt_cache_dir": "prompt_cache",
            "responses_dir": "artifacts/responses",
            "output_excel": "dist/output.xlsx",
        },
        "llm": {
            "provider": "openai_compatible",
            "base_url": "https://example.com/v1",
            "model": "custom-model",
            "api_key_env": "TEST_KEY",
            "headers": {"X-App-Name": "demand-conversion"},
            "options": {"endpoint": "responses"},
            "timeout_seconds": 45,
            "max_concurrency": 2,
            "max_retries": 4,
        },
        "batch": {"batch_size": 7},
        "flags": {"resume": True, "dry_run": True, "preview": True},
    }

    config = PipelineConfig.from_dict(payload, root=tmp_path)

    assert config.paths.docx == tmp_path / "custom.docx"
    assert config.paths.responses_dir == tmp_path / "artifacts/responses"
    assert config.paths.output_excel == tmp_path / "dist/output.xlsx"
    assert config.llm.provider == "openai_compatible"
    assert config.llm.base_url == "https://example.com/v1"
    assert config.llm.headers == {"X-App-Name": "demand-conversion"}
    assert config.llm.options == {"endpoint": "responses"}
    assert config.llm.max_concurrency == 2
    assert config.batch.batch_size == 7
    assert config.flags.resume is True
    assert config.flags.dry_run is True
    assert config.flags.preview is True


def test_pipeline_config_defaults_provider_for_legacy_payload(tmp_path: Path) -> None:
    payload = {
        "llm": {
            "base_url": "https://example.com/v1",
            "model": "legacy-model",
            "api_key_env": "TEST_KEY",
        }
    }

    config = PipelineConfig.from_dict(payload, root=tmp_path)

    assert config.llm.provider == "openai_compatible"
    assert config.llm.base_url == "https://example.com/v1"
    assert config.llm.options == {"endpoint": "responses"}
    assert config.paths.responses_dir == tmp_path / "artifacts/responses"
    assert config.paths.output_excel == tmp_path / "output/output_test_cases.xlsx"


def test_pipeline_config_requires_base_url_for_openai_compatible(tmp_path: Path) -> None:
    payload = {
        "llm": {
            "provider": "openai_compatible",
            "model": "test-model",
        }
    }

    with pytest.raises(ValueError, match="llm.base_url"):
        PipelineConfig.from_dict(payload, root=tmp_path)


def test_pipeline_config_requires_model_for_openai_compatible(tmp_path: Path) -> None:
    payload = {
        "llm": {
            "provider": "openai_compatible",
            "base_url": "https://example.com/v1",
        }
    }

    with pytest.raises(ValueError, match="llm.model"):
        PipelineConfig.from_dict(payload, root=tmp_path)


def test_pipeline_state_roundtrip(tmp_path: Path) -> None:
    artifacts_dir = tmp_path / "artifacts"
    state = PipelineState(artifacts_dir)
    state.record_stage("requirements", {"count": 3})

    assert state.stage_completed("requirements")

    reloaded = PipelineState(artifacts_dir)
    assert reloaded.stage_completed("requirements")
    manifest = reloaded.as_dict()
    assert manifest["requirements"].details["count"] == 3


def test_build_provider_selects_expected_classes() -> None:
    openai_config = LLMConfig(
        provider="openai_compatible",
        base_url="https://example.com/v1",
        model="test-model",
        api_key_env="TEST_API_KEY",
    )
    custom_config = LLMConfig(
        provider="custom_http",
        base_url="https://example.com/api",
        model="test-model",
        api_key_env=None,
    )
    anthropic_config = LLMConfig(
        provider="anthropic",
        model="claude-sonnet-4-5",
        api_key_env="ANTHROPIC_API_KEY",
    )

    assert isinstance(build_provider(openai_config, "dummy"), OpenAICompatibleProvider)
    assert isinstance(build_provider(custom_config, None), CustomHTTPProvider)
    assert build_provider(anthropic_config, "dummy").__class__.__name__ == "AnthropicProvider"

    with pytest.raises(ValueError, match="Unsupported LLM provider"):
        build_provider(LLMConfig(provider="unknown", model="x"), None)


def test_llm_client_returns_cache_when_available(tmp_path: Path) -> None:
    config = LLMConfig(
        provider="custom_http",
        base_url="https://example.com/api",
        model="test-model",
        api_key_env=None,
        timeout_seconds=5,
        max_concurrency=1,
        max_retries=1,
        retry_backoff_seconds=1,
    )
    client = LLMClient(config, tmp_path)

    cached = LLMResponse(
        chunk_id=1,
        raw_text="cached output",
        metadata={"source": "cache", "provider": "custom_http", "model": "test-model"},
    )
    client._save_cache(cached)  # type: ignore[attr-defined]

    class DummyProvider:
        async def generate(self, chunk_id: int, prompt: str) -> ProviderResult:
            raise AssertionError("HTTP request should not be made when cache exists")

        async def close(self) -> None:
            return None

    client.provider = DummyProvider()

    response = asyncio.run(client.send_prompt(1, "prompt text"))
    asyncio.run(client.close())

    assert response.raw_text == "cached output"
    assert response.metadata["source"] == "cache"


def test_llm_client_invalidates_cache_when_model_changes(tmp_path: Path) -> None:
    config = LLMConfig(provider="custom_http", base_url="https://example.com/api", model="new-model", api_key_env=None)
    client = LLMClient(config, tmp_path)
    stale_cache = LLMResponse(
        chunk_id=1,
        raw_text="old output",
        metadata={"provider": "custom_http", "model": "old-model", "usage": {}},
    )
    client._save_cache(stale_cache)  # type: ignore[attr-defined]

    assert client._load_cache(1) is None  # type: ignore[attr-defined]


def test_llm_client_send_batch_runs_preflight_before_requests(tmp_path: Path) -> None:
    config = LLMConfig(provider="custom_http", base_url="https://example.com/api", model="test-model", api_key_env=None)
    client = LLMClient(config, tmp_path)
    calls: list[int] = []

    class DummyProvider:
        async def generate(self, chunk_id: int, prompt: str) -> ProviderResult:
            calls.append(chunk_id)
            if chunk_id == 0:
                return ProviderResult(raw_text="OK", metadata={"provider": "custom_http", "model": "test-model", "request_endpoint": "/generate"})
            return ProviderResult(raw_text=f"result-{chunk_id}", metadata={"provider": "custom_http", "model": "test-model"})

        async def close(self) -> None:
            return None

    client.provider = DummyProvider()
    responses = asyncio.run(client.send_batch([LLMRequest(1, "a"), LLMRequest(2, "b")]))

    assert [response.raw_text for response in responses] == ["result-1", "result-2"]
    assert calls == [0, 1, 2]


def test_llm_client_send_batch_fails_when_preflight_returns_empty_text(tmp_path: Path) -> None:
    config = LLMConfig(provider="custom_http", base_url="https://example.com/api", model="test-model", api_key_env=None, max_retries=1)
    client = LLMClient(config, tmp_path)

    class DummyProvider:
        async def generate(self, chunk_id: int, prompt: str) -> ProviderResult:
            return ProviderResult(raw_text="", metadata={"provider": "custom_http", "model": "test-model", "request_endpoint": "/generate"})

        async def close(self) -> None:
            return None

    client.provider = DummyProvider()

    with pytest.raises(RuntimeError, match="preflight failed"):
        asyncio.run(client.send_batch([LLMRequest(1, "a")]))


def test_openai_provider_extracts_text_from_structured_message_content() -> None:
    config = LLMConfig(
        provider="openai_compatible",
        base_url="https://example.com/v1",
        model="gpt-5-codex",
        api_key_env="TEST_API_KEY",
        options={"endpoint": "chat_completions"},
        timeout_seconds=5,
        max_concurrency=1,
        max_retries=1,
        retry_backoff_seconds=1,
    )
    provider = OpenAICompatibleProvider(config, "dummy")

    completion = SimpleNamespace(
        id="resp_test",
        usage=SimpleNamespace(to_dict=lambda: {"completion_tokens": 12, "prompt_tokens": 5, "total_tokens": 17}),
        choices=[
            SimpleNamespace(
                message=SimpleNamespace(
                    content=[
                        {"type": "output_text", "text": " /模块A\n* 用例：登录成功 - 操作：输入正确账号密码并提交 - 预期：进入首页 "}
                    ]
                )
            )
        ],
    )

    class DummyCompletions:
        def create(self, *args, **kwargs):
            return completion

    class DummyChat:
        completions = DummyCompletions()

    class DummyClient:
        chat = DummyChat()

    provider.client = DummyClient()

    response = asyncio.run(provider.generate(1, "prompt text"))

    assert response.raw_text == "/模块A\n* 用例：登录成功 - 操作：输入正确账号密码并提交 - 预期：进入首页"
    assert response.metadata["text_source"] == "choices.message.content"
    assert response.metadata["provider"] == "openai_compatible"
    assert response.metadata["request_endpoint"] == "chat_completions"


def test_openai_provider_defaults_to_responses_endpoint() -> None:
    config = LLMConfig(
        provider="openai_compatible",
        base_url="https://example.com/v1",
        model="gpt-5-codex",
        api_key_env="TEST_API_KEY",
        timeout_seconds=5,
        max_concurrency=1,
        max_retries=1,
        retry_backoff_seconds=1,
    )
    provider = OpenAICompatibleProvider(config, "dummy")

    completion = SimpleNamespace(
        id="resp_default_responses",
        output_text="/模块D\n• 用例：默认走 responses - 操作：发起请求 - 预期：返回正文",
        usage=SimpleNamespace(to_dict=lambda: {"completion_tokens": 6, "prompt_tokens": 3, "total_tokens": 9}),
        choices=[],
    )

    class DummyResponses:
        def create(self, *args, **kwargs):
            return completion

    class DummyClient:
        responses = DummyResponses()

    provider.client = DummyClient()

    response = asyncio.run(provider.generate(1, "prompt text"))

    assert response.raw_text == "/模块D\n• 用例：默认走 responses - 操作：发起请求 - 预期：返回正文"
    assert response.metadata["request_endpoint"] == "responses"


def test_openai_provider_extracts_text_from_output_text_attribute() -> None:
    config = LLMConfig(
        provider="openai_compatible",
        base_url="https://example.com/v1",
        model="gpt-5-codex",
        api_key_env="TEST_API_KEY",
        options={"endpoint": "responses"},
        timeout_seconds=5,
        max_concurrency=1,
        max_retries=1,
        retry_backoff_seconds=1,
    )
    provider = OpenAICompatibleProvider(config, "dummy")

    completion = SimpleNamespace(
        id="resp_output_text",
        output_text="/模块C\n• 用例：直接取 output_text - 操作：发起请求 - 预期：拿到正文",
        usage=SimpleNamespace(to_dict=lambda: {"completion_tokens": 7, "prompt_tokens": 3, "total_tokens": 10}),
        choices=[],
    )

    class DummyResponses:
        def create(self, *args, **kwargs):
            return completion

    class DummyClient:
        responses = DummyResponses()

    provider.client = DummyClient()

    response = asyncio.run(provider.generate(1, "prompt text"))

    assert response.raw_text == "/模块C\n• 用例：直接取 output_text - 操作：发起请求 - 预期：拿到正文"
    assert response.metadata["text_source"] == "output_text"
    assert response.metadata["request_endpoint"] == "responses"


def test_llm_client_ignores_empty_cache_with_completion_tokens(tmp_path: Path) -> None:
    config = LLMConfig(
        provider="custom_http",
        base_url="https://example.com/api",
        model="gpt-5-codex",
        api_key_env=None,
        timeout_seconds=5,
        max_concurrency=1,
        max_retries=1,
        retry_backoff_seconds=1,
    )
    client = LLMClient(config, tmp_path)
    stale_cache = LLMResponse(
        chunk_id=1,
        raw_text="",
        metadata={"provider": "custom_http", "model": "gpt-5-codex", "usage": {"completion_tokens": 9, "prompt_tokens": 4, "total_tokens": 13}},
    )
    client._save_cache(stale_cache)  # type: ignore[attr-defined]

    class DummyProvider:
        def __init__(self) -> None:
            self.calls = 0

        async def generate(self, chunk_id: int, prompt: str) -> ProviderResult:
            self.calls += 1
            return ProviderResult(
                raw_text="/模块B\n• 用例：刷新缓存 - 操作：重新请求 - 预期：写入正文",
                metadata={"provider": "custom_http", "model": "gpt-5-codex", "text_source": "choices.message.content", "usage": {"completion_tokens": 8}},
            )

        async def close(self) -> None:
            return None

    provider = DummyProvider()
    client.provider = provider

    response = asyncio.run(client.send_prompt(1, "prompt text"))

    assert provider.calls == 1
    assert response.raw_text == "/模块B\n• 用例：刷新缓存 - 操作：重新请求 - 预期：写入正文"
    assert response.metadata["text_source"] == "choices.message.content"


def test_custom_http_provider_extracts_text_and_substitutes_headers() -> None:
    config = LLMConfig(
        provider="custom_http",
        base_url="https://example.com/api",
        model="testcase-v1",
        api_key_env="CUSTOM_LLM_KEY",
        headers={"Authorization": "Bearer ${API_KEY}", "X-App-Name": "demand-conversion"},
        options={
            "method": "POST",
            "path": "/generate",
            "request_format": "prompt_model",
            "response_text_path": "data.output",
            "response_usage_path": "meta.usage",
        },
    )
    provider = CustomHTTPProvider(config, "secret-token")

    assert provider._build_headers()["Authorization"] == "Bearer secret-token"
    assert provider._build_request_body("hello") == {"model": "testcase-v1", "prompt": "hello"}

    provider._perform_request = lambda prompt: (
        200,
        {
            "id": "custom-1",
            "data": {"output": "/模块D\n• 用例：自定义协议 - 操作：发请求 - 预期：提取正文"},
            "meta": {"usage": {"completion_tokens": 6, "total_tokens": 9}},
        },
    )

    response = asyncio.run(provider.generate(1, "hello"))

    assert response.raw_text == "/模块D\n• 用例：自定义协议 - 操作：发请求 - 预期：提取正文"
    assert response.metadata["text_source"] == "data.output"
    assert response.metadata["usage"]["completion_tokens"] == 6
    assert response.metadata["status_code"] == 200


def test_custom_http_provider_supports_messages_model_request_format() -> None:
    config = LLMConfig(
        provider="custom_http",
        base_url="https://example.com/api",
        model="testcase-v1",
        api_key_env=None,
        options={"request_format": "messages_model"},
    )
    provider = CustomHTTPProvider(config, None)

    assert provider._build_request_body("hello") == {
        "model": "testcase-v1",
        "messages": [{"role": "user", "content": "hello"}],
    }


def test_pipeline_stages_load_responses_invalidates_blank_artifacts_with_completion_tokens(tmp_path: Path) -> None:
    responses_dir = tmp_path / "responses"
    responses_dir.mkdir(parents=True)
    (responses_dir / "chunk_001.txt").write_text("", encoding="utf-8")

    prompt_cache_dir = tmp_path / "prompt_cache"
    prompt_cache_dir.mkdir()
    (prompt_cache_dir / "chunk_001.json").write_text(
        '{"raw_text": "", "metadata": {"provider": "openai_compatible", "model": "gpt-5-codex", "usage": {"completion_tokens": 5, "prompt_tokens": 2, "total_tokens": 7}}}',
        encoding="utf-8",
    )

    artifacts_dir = tmp_path / "artifacts"
    config = PipelineConfig(
        paths=PathsConfig(
            docx=tmp_path / "source.docx",
            artifacts_dir=artifacts_dir,
            prompt_cache_dir=prompt_cache_dir,
            responses_dir=responses_dir,
            output_excel=tmp_path / "output.xlsx",
        ),
        llm=LLMConfig(base_url="https://example.com/v1", model="gpt-5-codex", api_key_env="TEST_API_KEY"),
        batch=BatchConfig(batch_size=1),
        flags=ExecutionFlags(),
    )
    state = PipelineState(artifacts_dir)

    stages = PipelineStages(config, state)

    assert stages.load_responses() is None


def test_pipeline_stages_call_llm_writes_to_configured_responses_dir(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    import pipeline.stages as stages_module

    responses_dir = tmp_path / "custom-responses"
    config = PipelineConfig(
        paths=PathsConfig(
            docx=tmp_path / "source.docx",
            artifacts_dir=tmp_path / "artifacts",
            prompt_cache_dir=tmp_path / "prompt_cache",
            responses_dir=responses_dir,
            output_excel=tmp_path / "output.xlsx",
        ),
        llm=LLMConfig(provider="custom_http", base_url="https://example.com/api", model="gpt-5-codex", api_key_env=None),
        batch=BatchConfig(batch_size=1),
        flags=ExecutionFlags(),
    )
    state = PipelineState(config.paths.artifacts_dir)

    class DummyLLMClient:
        def __init__(self, config, cache_dir):
            self.config = config
            self.cache_dir = cache_dir

        async def send_batch(self, requests):
            return [LLMResponse(chunk_id=1, raw_text="ok", metadata={"provider": "custom_http"})]

        async def close(self):
            return None

    monkeypatch.setattr(stages_module, "LLMClient", DummyLLMClient)

    stages = PipelineStages(config, state)
    responses = asyncio.run(stages.call_llm(["prompt text"]))

    assert responses == ["ok"]
    assert (responses_dir / "chunk_001.txt").read_text(encoding="utf-8") == "ok"
    assert state.stage_completed("llm") is True


def test_dry_run_ignores_stale_responses(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    import pipeline.stages as stages_module

    responses_dir = tmp_path / "responses"
    responses_dir.mkdir()
    (responses_dir / "chunk_001.txt").write_text(
        "/旧模块\n• 用例：旧响应 - 操作：旧操作 - 预期：旧预期",
        encoding="utf-8",
    )

    config = PipelineConfig(
        paths=PathsConfig(
            docx=tmp_path / "source.docx",
            artifacts_dir=tmp_path / "artifacts",
            prompt_cache_dir=tmp_path / "prompt_cache",
            responses_dir=responses_dir,
            output_excel=tmp_path / "output.xlsx",
        ),
        llm=LLMConfig(provider="custom_http", base_url="https://example.com/api", model="dry-run", api_key_env=None),
        batch=BatchConfig(batch_size=1),
        flags=ExecutionFlags(dry_run=True),
    )
    state = PipelineState(config.paths.artifacts_dir)
    state.record_stage("llm", {"count": 1})
    state.record_stage("cases", {"count": 1})
    state.record_stage("excel", {"rows": 1})
    monkeypatch.setattr(
        stages_module,
        "_parse_requirement_tables",
        lambda _: [
            RequirementBlock(
                module="新模块",
                title="新需求",
                description="新描述",
            )
        ],
    )

    asyncio.run(run_pipeline(config, state))

    assert state.stage_completed("requirements") is True
    assert state.stage_completed("prompts") is True
    assert state.stage_completed("llm") is False
    assert state.stage_completed("cases") is False
    assert state.stage_completed("excel") is False
    assert not (config.paths.artifacts_dir / "parsed_cases.json").exists()


def test_llm_output_parser_accepts_asterisk_bullets() -> None:
    text = "/模块A\n* 用例：登录成功 - 操作：输入正确账号密码并提交 - 预期：进入首页"

    cases = LLMOutputParser(text).parse()

    assert len(cases) == 1
    assert cases[0].module == "/模块A"
    assert cases[0].name == "登录成功"
    assert cases[0].operation == "输入正确账号密码并提交"
    assert cases[0].expected == "进入首页"


def test_validate_docx_prints_on_gbk_console(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    docx_path = tmp_path / "requirements.docx"
    document = Document()
    document.add_paragraph("1. 模块")
    document.add_paragraph("需求描述")
    document.save(docx_path)

    stdout_buffer = io.BytesIO()
    gbk_stdout = io.TextIOWrapper(stdout_buffer, encoding="gbk", errors="strict")
    monkeypatch.setattr(sys, "stdout", gbk_stdout)

    validate_docx(
        docx_path,
        tmp_path / "requirements.json",
        tmp_path / "requirements_preview.md",
        "utf-8",
    )
