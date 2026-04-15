"""CLI helper to extract structured requirement rows from the Word source.

Usage::

    conda run -n demand_conversion python requirement_extractor.py \
        --docx source/wordv1.3.7.docx --format json --output requirements.json

Run with ``--format markdown`` to emit a Markdown summary similar to
``requirements.md``.
"""
from __future__ import annotations

import argparse
import json
import logging
from collections import OrderedDict
from pathlib import Path
from typing import Iterable, List, MutableMapping, Sequence

from docx import Document
from docx.text.paragraph import Paragraph

from requirement_models import RequirementBlock
from requirement_parser import parse_requirement_tables

DEFAULT_DOCX = Path("source/wordv1.3.7.docx")
DEFAULT_PREVIEW = Path("tmp_table_preview.txt")


class OutputFormat:
    JSON = "json"
    MARKDOWN = "markdown"
    PROMPT_JSON = "prompt_json"


class PreviewFormat:
    PRETTY = "pretty"
    RAW = "raw"


def main() -> None:
    args = _parse_args()
    log_level = getattr(logging, args.log_level.upper(), logging.INFO)
    logging.basicConfig(level=log_level, format="%(levelname)s: %(message)s")

    docx_path = Path(args.docx)
    if not docx_path.exists():
        raise SystemExit(f"DOCX file not found: {docx_path}")

    logging.info("Parsing requirements from %s", docx_path)
    blocks = parse_requirement_tables(docx_path)
    if not blocks:
        raise SystemExit("No requirement blocks were parsed from the document.")

    logging.info("Parsed %d requirement rows", len(blocks))

    if args.format == OutputFormat.MARKDOWN:
        content = _render_markdown(blocks)
    elif args.format == OutputFormat.PROMPT_JSON:
        content = _render_prompt_json(blocks)
    else:
        content = _render_json(blocks)

    _write_output(content, args.output, args.encoding)
    logging.info("Wrote %s output", args.format)

    if args.preview:
        preview_path = Path(args.preview)
        logging.info(
            "Writing preview to %s (first %d blocks, %d paragraphs)",
            preview_path,
            args.preview_limit,
            args.preview_paragraphs,
        )
        _write_preview(
            docx_path,
            preview_path,
            blocks,
            block_limit=args.preview_limit,
            paragraph_limit=args.preview_paragraphs,
            encoding=args.encoding,
            format=args.preview_format,
        )
        logging.info("Preview format: %s", args.preview_format)
        logging.info("Preview file written: %s", preview_path)


def _parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description=(
            "Extract requirement blocks from a Word doc and emit JSON/Markdown for "
            "downstream automation."
        )
    )
    parser.add_argument(
        "--docx",
        default=str(DEFAULT_DOCX),
        help="Path to the Word requirements file.",
    )
    parser.add_argument(
        "--format",
        choices=(OutputFormat.JSON, OutputFormat.MARKDOWN, OutputFormat.PROMPT_JSON),
        default=OutputFormat.JSON,
        help="Choose json, prompt_json or markdown output (default: json).",
    )
    parser.add_argument(
        "--output",
        help="Destination file. Omit to print to stdout.",
    )
    parser.add_argument(
        "--encoding",
        default="utf-8",
        help="Encoding for written files (default: utf-8).",
    )
    parser.add_argument(
        "--log-level",
        default="INFO",
        help="Logging level (DEBUG, INFO, WARNING, ERROR).",
    )
    parser.add_argument(
        "--preview",
        nargs="?",
        const=str(DEFAULT_PREVIEW),
        help="Optional path to write a human-friendly preview (default file when flag is set).",
    )
    parser.add_argument(
        "--preview-format",
        choices=(PreviewFormat.PRETTY, PreviewFormat.RAW),
        default=PreviewFormat.PRETTY,
        help="Preview style: pretty summary or raw JSON block dump (default: pretty).",
    )
    parser.add_argument(
        "--preview-limit",
        type=int,
        default=10,
        help="How many parsed blocks to include in the preview (default: 10).",
    )
    parser.add_argument(
        "--preview-paragraphs",
        type=int,
        default=8,
        help="How many raw paragraphs (with styles) to include in the preview (default: 8).",
    )
    return parser.parse_args()


def _render_json(blocks: Sequence[RequirementBlock]) -> str:
    return json.dumps([block.as_dict() for block in blocks], ensure_ascii=False, indent=2) + "\n"


def _render_prompt_json(blocks: Sequence[RequirementBlock]) -> str:
    payload = {
        "requirements": [block.as_dict() for block in blocks],
        "metadata": {
            "total": len(blocks),
        },
    }
    return json.dumps(payload, ensure_ascii=False, indent=2) + "\n"


def _render_markdown(blocks: Sequence[RequirementBlock]) -> str:
    grouped = _group_by_module(blocks)
    lines: List[str] = ["# Requirement Summary"]
    for module, module_blocks in grouped.items():
        lines.append("")
        lines.append(f"## {module}")
        for idx, block in enumerate(module_blocks, start=1):
            lines.append("")
            lines.append(f"### {idx}. {block.title}")
            lines.append("")
            description = block.description.strip() or "(无描述)"
            lines.append(description)
            if block.notes:
                lines.append("")
                lines.append("示例 / 备注：")
                lines.append(block.notes.strip())
    return "\n".join(lines).strip() + "\n"


def _group_by_module(blocks: Sequence[RequirementBlock]) -> MutableMapping[str, List[RequirementBlock]]:
    grouped: MutableMapping[str, List[RequirementBlock]] = OrderedDict()
    for block in blocks:
        grouped.setdefault(block.module, []).append(block)
    return grouped


def _write_output(content: str, output_path: str | None, encoding: str) -> None:
    if output_path:
        destination = Path(output_path)
        destination.parent.mkdir(parents=True, exist_ok=True)
        destination.write_text(content, encoding=encoding)
    else:
        print(content)


def _write_preview(
    docx_path: Path,
    preview_path: Path,
    blocks: Sequence[RequirementBlock],
    *,
    block_limit: int,
    paragraph_limit: int,
    encoding: str,
    format: str,
) -> None:
    block_limit = max(block_limit, 0)
    paragraph_limit = max(paragraph_limit, 0)
    document = Document(docx_path)

    if format == PreviewFormat.RAW:
        content = json.dumps(
            {
                "requirements": [block.as_dict() for block in blocks[:block_limit]],
                "paragraphs": [
                    {
                        "style": _safe_style_name(paragraph),
                        "text": paragraph.text,
                    }
                    for paragraph in document.paragraphs[:paragraph_limit]
                ],
            },
            ensure_ascii=False,
            indent=2,
        )
        preview_path.parent.mkdir(parents=True, exist_ok=True)
        preview_path.write_text(content + "\n", encoding=encoding)
        return

    lines: List[str] = [f"Preview generated from {docx_path}"]
    lines.append("")
    lines.append("== Parsed Requirement Blocks ==")
    if not blocks:
        lines.append("(no blocks parsed)")
    else:
        for idx, block in enumerate(blocks[:block_limit], start=1):
            lines.append(f"[{idx}] {block.module} -> {block.title}")
            snippet = block.description.strip().replace("\n", " ")
            lines.append(f"    描述: {snippet[:180]}{'…' if len(snippet) > 180 else ''}")
            if block.segments:
                lines.append(f"    Segments: {len(block.segments)}")
            if block.block_id:
                lines.append(f"    Block ID: {block.block_id}")
            if block.notes:
                notes_snippet = block.notes.strip().replace("\n", " ")
                lines.append(
                    f"    备注: {notes_snippet[:120]}{'…' if len(notes_snippet) > 120 else ''}"
                )

    lines.append("")
    lines.append("== Paragraph Style Snapshot ==")
    paragraphs: Iterable[Paragraph] = document.paragraphs[:paragraph_limit]
    if not paragraphs:
        lines.append("(no paragraphs found)")
    else:
        for idx, paragraph in enumerate(paragraphs, start=1):
            style_name = _safe_style_name(paragraph)
            text = paragraph.text.strip().replace("\n", " ")
            lines.append(f"Paragraph {idx} [{style_name}]: {text}")

    preview_path.parent.mkdir(parents=True, exist_ok=True)
    preview_path.write_text("\n".join(lines).strip() + "\n", encoding=encoding)


def _safe_style_name(paragraph: Paragraph) -> str:
    try:
        style = paragraph.style
        if style is None:
            return "(no style)"
        name = getattr(style, "name", None)
        return name or "(unnamed style)"
    except AttributeError:
        return "(missing style)"


if __name__ == "__main__":
    main()
